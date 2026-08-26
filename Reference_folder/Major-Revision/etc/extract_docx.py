import sys
import docx

def extract(path, out_path):
    d = docx.Document(path)
    lines = []
    for i, para in enumerate(d.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text
        if text.strip():
            lines.append(f"[{style}] {text}")
    # also extract tables
    for ti, table in enumerate(d.tables):
        lines.append(f"\n=== TABLE {ti} ===")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Wrote {len(lines)} lines to {out_path}")

if __name__ == "__main__":
    src = sys.argv[1]
    out = sys.argv[2]
    extract(src, out)
