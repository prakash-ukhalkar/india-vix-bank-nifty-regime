"""
fix_references.py
1. Insert GATE-GNN (Fofanah et al., 2024) and EATSA-GNN (Fofanah & Leigh, 2025)
   into the reference list, between GBK-GNN (P[187]) and PIMPC-GNN (P[188]).
2. Update the body-text citation in P[16] to include the new references.
"""

import copy
from lxml import etree
import docx
from docx.oxml.ns import qn

PATH = r"f:\MLSAPU\PhD-SPPU\India-VIX-Major-Revision\Major-Revision\Latest_Main-Manuscript_UPDATED_2026-07-24_tables_fixed.docx"

d = docx.Document(PATH)
paras = d.paragraphs

# -----------------------------------------------------------------------
# 1. Locate the anchor paragraph (GBK-GNN, Du et al. 2021)
# -----------------------------------------------------------------------
anchor_idx = None
for i, p in enumerate(paras):
    if "GBK-GNN" in p.text and "Du" in p.text:
        anchor_idx = i
        break

if anchor_idx is None:
    raise RuntimeError("Cannot find GBK-GNN reference paragraph")

print(f"Anchor paragraph (GBK-GNN): P[{anchor_idx}]")

anchor_para = paras[anchor_idx]

# -----------------------------------------------------------------------
# 2. Build two new reference paragraphs as XML, copying style from anchor
# -----------------------------------------------------------------------
GATE_GNN_TEXT = (
    "Fofanah, A. J., Chen, D., Wen, L., & Zhang, S. (2024). "
    "Addressing imbalance in graph datasets: Introducing GATE-GNN with "
    "graph ensemble weight attention and transfer learning for enhanced node "
    "classification. Expert Systems with Applications, 255, 124602. "
    "https://doi.org/10.1016/j.eswa.2024.124602"
)

EATSA_GNN_TEXT = (
    "Fofanah, A. J., & Leigh, A. O. (2025). "
    "EATSA-GNN: Edge-Aware and Two-Stage attention for enhancing graph neural "
    "networks based on teacher\u2013student mechanisms for graph node classification. "
    "Neurocomputing, 612, 128686. "
    "https://doi.org/10.1016/j.neucom.2024.128686"
)


def make_para_xml(template_para, text):
    """Clone the template paragraph's XML and replace its run text."""
    new_xml = copy.deepcopy(template_para._element)
    # Remove all existing runs
    for r in new_xml.findall(qn("w:r")):
        new_xml.remove(r)
    # Add a single run with the new text
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_el = etree.SubElement(new_xml, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return new_xml


gate_xml  = make_para_xml(anchor_para, GATE_GNN_TEXT)
eatsa_xml = make_para_xml(anchor_para, EATSA_GNN_TEXT)

# Insert AFTER the anchor paragraph (GBK-GNN):
# order will be: GBK-GNN → GATE-GNN → EATSA-GNN → PIMPC-GNN
anchor_el = anchor_para._element
parent    = anchor_el.getparent()
idx_in_parent = list(parent).index(anchor_el)

# Insert EATSA first (so it ends up after GATE when both are inserted after anchor)
parent.insert(idx_in_parent + 1, eatsa_xml)
parent.insert(idx_in_parent + 1, gate_xml)

print("Inserted GATE-GNN and EATSA-GNN references after GBK-GNN.")

# -----------------------------------------------------------------------
# 3. Update body-text citation in P[16]
#    Replace the old "(Du et al., 2021; Fofanah et al., 2026)"
#    with expanded citation including 2024 and 2025 papers
# -----------------------------------------------------------------------
OLD_CITE = "(Du et al., 2021; Fofanah et al., 2026)"
NEW_CITE = "(Du et al., 2021; Fofanah et al., 2024; Fofanah & Leigh, 2025; Fofanah et al., 2026)"

body_para = None
for p in d.paragraphs:
    if OLD_CITE in p.text:
        body_para = p
        break

if body_para is None:
    print("WARNING: Could not find the citation text to update in body paragraph.")
else:
    # Walk all runs in the paragraph, find and replace the citation text
    full_text = body_para.text
    if OLD_CITE in full_text:
        # Rebuild runs: find which run(s) contain the citation
        # Simplest: concatenate all run texts, do replacement, rewrite into first run
        runs = body_para.runs
        combined = "".join(r.text for r in runs)
        new_combined = combined.replace(OLD_CITE, NEW_CITE, 1)
        # Write back to first run, blank the rest
        if runs:
            runs[0].text = new_combined
            for r in runs[1:]:
                r.text = ""
        print(f"Updated citation in body paragraph: {body_para.text[-80:]}")
    else:
        print("Citation text not found in runs (may span runs). Attempting XML walk.")
        # Fallback: walk XML text nodes
        para_el = body_para._element
        texts = para_el.findall(".//" + qn("w:t"))
        concat = "".join(t.text or "" for t in texts)
        if OLD_CITE in concat:
            new_concat = concat.replace(OLD_CITE, NEW_CITE, 1)
            # Wipe all t elements and put new text in first
            for t in texts:
                t.text = ""
            if texts:
                texts[0].text = new_concat
            print("Citation updated via XML fallback.")
        else:
            print("WARNING: OLD_CITE not found even in XML concatenation.")

# -----------------------------------------------------------------------
# 4. Save
# -----------------------------------------------------------------------
d.save(PATH)
print(f"\nFile saved: {PATH}")

# Quick verification
d2 = docx.Document(PATH)
print("\n=== Verification: last 10 reference paragraphs ===")
for i, p in enumerate(d2.paragraphs):
    if any(k in p.text for k in ["GBK-GNN", "GATE-GNN", "EATSA-GNN", "PIMPC-GNN"]):
        print(f"  P[{i}]: {p.text[:120]}")
print("\n=== Citation in body (P[16] area) ===")
for i, p in enumerate(d2.paragraphs):
    if "GATE-GNN" in p.text and "supervised" in p.text:
        print(f"  P[{i}]: ...{p.text[-150:]}")
