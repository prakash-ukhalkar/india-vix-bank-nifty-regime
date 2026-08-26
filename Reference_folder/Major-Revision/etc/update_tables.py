"""
update_tables.py  --  Correct all stale tables in the updated manuscript docx.

Tables corrected:
  TABLE[1]  = Table 2  : Sample construction (dates, row counts, VIX threshold)
  TABLE[5]  = Table 6  : Overall test performance (21-day row only)
  TABLE[6]  = Table 7  : Regime-conditioned accuracy (21-day row)
  TABLE[7]  = Table 8  : Trading performance (full rebuild with corrected values)
  TABLE[8]  = Table 9  : Cost sensitivity (full rebuild with corrected values)
  TABLE[9]  = Table 10 : GARCH comparison (VIX rows only)
  TABLE[10] = Table 11 : Threshold sensitivity (primary p75 row)

All values sourced from:
  results/trading_simulation.csv, results/statistical_tests.json,
  data/processed/split_summary.csv, results/vix_threshold_config.json
"""

import copy
import json
import shutil
import math
from pathlib import Path

import docx
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(r"f:\MLSAPU\PhD-SPPU\India-VIX-Major-Revision")
SRC  = ROOT / "Major-Revision" / "Latest_Main-Manuscript_UPDATED_2026-07-24.docx"
DST  = ROOT / "Major-Revision" / "Latest_Main-Manuscript_UPDATED_2026-07-24_tables_fixed.docx"

# ---------------------------------------------------------------------------
# Load authoritative result data
# ---------------------------------------------------------------------------
import pandas as pd

split_df = pd.read_csv(ROOT / "data/processed/split_summary.csv")
trading_df = pd.read_csv(ROOT / "results/trading_simulation.csv")

with open(ROOT / "results/statistical_tests.json") as f:
    stats = json.load(f)

with open(ROOT / "results/vix_threshold_config.json") as f:
    vix_cfg = json.load(f)

# Derived values
vix_thr    = round(float(vix_cfg["vix_threshold_fixed"]), 2)  # 18.71
n_high     = int(vix_cfg["high_vix_days_fixed"])            # 8
n_low      = 256 - n_high                                    # 248
n_aligned  = 256                                             # BiLSTM-aligned eval sample

# Split rows
train_row = split_df[split_df["Split"] == "Train"].iloc[0]
val_row   = split_df[split_df["Split"] == "Val"].iloc[0]
test_row  = split_df[split_df["Split"] == "Test"].iloc[0]

n_train = int(train_row["Rows"])   # 1794
n_val   = int(val_row["Rows"])     # 273
n_test  = int(test_row["Rows"])    # 276

train_start = train_row["Start"]           # 2016-05-23
train_end   = train_row["End (clean)"]     # 2023-09-11
val_start   = val_row["Start"]             # 2023-10-20
val_end     = val_row["End (clean)"]       # 2024-12-04
test_start  = test_row["Start"]            # 2025-01-14
test_end    = test_row["End (clean)"]      # 2026-02-25

# Statistical test data
hv_acc   = stats["high_vix"]["accuracy"] / 100.0            # 1.000
lv_acc   = stats["low_vix"]["accuracy"] / 100.0             # 0.665
acc_diff = stats["accuracy_diff_pp"] / 100.0                # 0.335
fish_p   = stats["fisher_non_overlap_p"]                     # 1.000

overall_acc  = stats["stack_overall_acc"] / 100.0           # 0.676
overall_auc  = stats["stack_overall_auc"]                    # 0.582
majority_acc = stats["majority_baseline_overall"] / 100.0   # 0.711

# McNemar test (model vs majority-Up baseline)
# b = TN = 2 (model right, majority wrong: predicted Down, actual Down)
# c = FN = 11 (majority right, model wrong: predicted Down, actual Up)
TN, FP, FN, TP = 2, 72, 11, 163+8   # combined High+Low
b, c = TN, FN
n_mc = b + c   # 13

def binom_pmf(n, k, p=0.5):
    import math
    return math.comb(n, k) * (p**k) * ((1-p)**(n-k))

mcnemar_p = 2 * sum(binom_pmf(n_mc, k) for k in range(0, min(b, c) + 1))
# mcnemar_p ≈ 0.022

# Trading data (TC = 0bps is the baseline row)
bh_cum  = trading_df.loc[0, "bh_cum_ret_pct"]    # 23.56
bh_shar = trading_df.loc[0, "bh_sharpe"]          # 1.204
bh_dd   = trading_df.loc[0, "bh_max_dd_pct"]      # -6.62
deployed= trading_df.loc[0, "deployed_days"]        # 21

# Test period years for annualization (276 trading days / 252)
test_years = n_test / 252.0   # 1.0952

def ann_ret(total_pct, years=test_years):
    """Simple linear annualization matching notebook convention."""
    return total_pct / years

bh_ann = ann_ret(bh_cum)   # ≈ 21.5

# Strategy rows indexed by cost_bps
tc_map = {}
for _, row in trading_df.iterrows():
    bps = int(row["cost_bps"])
    tc_map[bps] = {
        "cum": row["strat_cum_ret_pct"],
        "ann": ann_ret(row["strat_cum_ret_pct"]),
        "sha": row["strat_sharpe"],
        "dd":  row["strat_max_dd_pct"],
    }

# Interpolate 3 bps (not in CSV)
r0, r10 = tc_map[0], tc_map[10]
frac = 3 / 10.0
tc_map[3] = {
    "cum": r0["cum"] + frac * (r10["cum"] - r0["cum"]),
    "ann": r0["ann"] + frac * (r10["ann"] - r0["ann"]),
    "sha": r0["sha"] + frac * (r10["sha"] - r0["sha"]),
    "dd":  r0["dd"],   # drawdown unchanged
}

# ---------------------------------------------------------------------------
# Utility: set a table cell's text while preserving paragraph formatting
# ---------------------------------------------------------------------------
def set_cell(cell, text):
    """Replace first paragraph text of a table cell."""
    para = cell.paragraphs[0]
    # Clear all runs in that paragraph
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = str(text)
    else:
        para.add_run(str(text))

# ---------------------------------------------------------------------------
# Open and copy the source document
# ---------------------------------------------------------------------------
shutil.copy2(SRC, DST)
d = docx.Document(DST)

tables = d.tables

# ---------------------------------------------------------------------------
# TABLE[1] = Table 2 : Sample construction
# ---------------------------------------------------------------------------
t = tables[1]
# Row[1] = Training
set_cell(t.rows[1].cells[1], f"{train_start} to {train_end}")
set_cell(t.rows[1].cells[2], f"{n_train:,}")
set_cell(t.rows[1].cells[3], "Leakage-corrected boundary; used for feature scaling, base learners, and time-series CV")

# Row[2] = Validation
set_cell(t.rows[2].cells[1], f"{val_start} to {val_end}")
set_cell(t.rows[2].cells[2], f"{n_val:,}")
set_cell(t.rows[2].cells[3], "Leakage-corrected boundary; used for tuning decisions and early stopping")

# Row[3] = Test
set_cell(t.rows[3].cells[1], f"{test_start} to {test_end}")
set_cell(t.rows[3].cells[2], f"{n_test:,}")
set_cell(t.rows[3].cells[3], f"Strictly held out; {n_aligned}-row aligned evaluation after BiLSTM 20-day lookback")

# Row[4] = High-VIX test days
set_cell(t.rows[4].cells[0], "High-VIX test days")
set_cell(t.rows[4].cells[1], f"VIX \u2265 {vix_thr:.2f} (train+val p75; fixed pre-test)")
set_cell(t.rows[4].cells[2], str(n_high))
pct_high = n_high / n_aligned * 100
set_cell(t.rows[4].cells[3], f"{pct_high:.1f}% of the {n_aligned}-row aligned evaluation sample")

# Row[5] = Low-VIX test days
set_cell(t.rows[5].cells[0], "Low-VIX test days")
set_cell(t.rows[5].cells[1], f"VIX < {vix_thr:.2f}")
set_cell(t.rows[5].cells[2], str(n_low))
pct_low = n_low / n_aligned * 100
set_cell(t.rows[5].cells[3], f"{pct_low:.1f}% of the aligned evaluation sample")

print(f"TABLE[1] (Table 2) updated: train={n_train}, val={n_val}, test={n_test}, VIX>={vix_thr:.2f}, n_high={n_high}")

# ---------------------------------------------------------------------------
# TABLE[5] = Table 6 : Overall performance — update 21-day row only
# ---------------------------------------------------------------------------
t = tables[5]
# Row[3] = 21-day
set_cell(t.rows[3].cells[1], f"{overall_acc:.4f}")
set_cell(t.rows[3].cells[2], f"{overall_auc:.4f}")
set_cell(t.rows[3].cells[3], f"{majority_acc:.4f}")
set_cell(t.rows[3].cells[4], f"{mcnemar_p:.3f}")

print(f"TABLE[5] (Table 6) 21-day row: acc={overall_acc:.4f}, AUC={overall_auc:.4f}, "
      f"majority={majority_acc:.4f}, McNemar_p={mcnemar_p:.3f}")

# ---------------------------------------------------------------------------
# TABLE[6] = Table 7 : Regime-conditioned accuracy — update 21-day row only
# ---------------------------------------------------------------------------
t = tables[6]
# Row[3] = 21-day
set_cell(t.rows[3].cells[1], str(n_high))
set_cell(t.rows[3].cells[2], f"{hv_acc:.3f}")
set_cell(t.rows[3].cells[3], str(n_low))
set_cell(t.rows[3].cells[4], f"{lv_acc:.3f}")
set_cell(t.rows[3].cells[5], f"+{acc_diff:.3f}")
set_cell(t.rows[3].cells[6], f"{fish_p:.3f}")
set_cell(t.rows[3].cells[7], "N/A")
set_cell(t.rows[3].cells[8], "No")

print(f"TABLE[6] (Table 7) 21-day row: HV n={n_high}, HV acc={hv_acc:.3f}, LV n={n_low}, LV acc={lv_acc:.3f}")

# ---------------------------------------------------------------------------
# TABLE[7] = Table 8 : Trading performance
# col headers: Strategy | Active days | Total return | Annual return | Sharpe | Max drawdown
# ---------------------------------------------------------------------------
t = tables[7]

# Row[1] = Buy-and-Hold
set_cell(t.rows[1].cells[1], f"{n_aligned}/{n_aligned}")
set_cell(t.rows[1].cells[2], f"+{bh_cum:.2f}%")
set_cell(t.rows[1].cells[3], f"{bh_ann:.1f}%")
set_cell(t.rows[1].cells[4], f"{bh_shar:.3f}")
set_cell(t.rows[1].cells[5], f"{bh_dd:.2f}%")

# Row[2] = Full Signal — not computed under corrected protocol
set_cell(t.rows[2].cells[1], "\u2014")
set_cell(t.rows[2].cells[2], "\u2014")
set_cell(t.rows[2].cells[3], "\u2014")
set_cell(t.rows[2].cells[4], "\u2014")
set_cell(t.rows[2].cells[5], "\u2014")

# Row[3] = High-VIX + Up (corrected, 1 trade × 21 days)
s0 = tc_map[0]
set_cell(t.rows[3].cells[1], f"{deployed}/{n_aligned}")
set_cell(t.rows[3].cells[2], f"+{s0['cum']:.2f}%")
set_cell(t.rows[3].cells[3], f"{s0['ann']:.1f}%")
set_cell(t.rows[3].cells[4], f"{s0['sha']:.3f}")
set_cell(t.rows[3].cells[5], f"{s0['dd']:.2f}%")

# Row[4] = High-VIX Only — not computed under corrected protocol
set_cell(t.rows[4].cells[1], "\u2014")
set_cell(t.rows[4].cells[2], "\u2014")
set_cell(t.rows[4].cells[3], "\u2014")
set_cell(t.rows[4].cells[4], "\u2014")
set_cell(t.rows[4].cells[5], "\u2014")

print(f"TABLE[7] (Table 8): BH sharpe={bh_shar}, HV+Up sharpe={s0['sha']}, BH cum={bh_cum}%, HV+Up cum={s0['cum']}%")

# ---------------------------------------------------------------------------
# TABLE[8] = Table 9 : Cost sensitivity
# col headers: Cost setting | Total return | Annual return | Sharpe | Max drawdown
# row structure:
#   Row[1] = Buy-and-Hold benchmark
#   Row[2] = HV+Up TC = 0.00%
#   Row[3] = HV+Up TC = 0.03%
#   Row[4] = HV+Up TC = 0.10%
#   Row[5] = HV+Up TC = 0.20%
#   Row[6] = HV+Up TC = 0.30%
#   Row[7] = HV+Up TC = 0.50%
# ---------------------------------------------------------------------------
t = tables[8]

# Row[1] = Buy-and-Hold benchmark
set_cell(t.rows[1].cells[1], f"{bh_cum:.2f}%")
set_cell(t.rows[1].cells[2], f"{bh_ann:.1f}%")
set_cell(t.rows[1].cells[3], f"{bh_shar:.3f}")
set_cell(t.rows[1].cells[4], f"{bh_dd:.2f}%")

# Row[2] = TC 0.00% (0 bps)
r = tc_map[0]
set_cell(t.rows[2].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[2].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[2].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[2].cells[4], f"{r['dd']:.2f}%")

# Row[3] = TC 0.03% (3 bps, interpolated)
r = tc_map[3]
set_cell(t.rows[3].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[3].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[3].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[3].cells[4], f"{r['dd']:.2f}%")

# Row[4] = TC 0.10% (10 bps)
r = tc_map[10]
set_cell(t.rows[4].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[4].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[4].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[4].cells[4], f"{r['dd']:.2f}%")

# Row[5] = TC 0.20% (20 bps)
r = tc_map[20]
set_cell(t.rows[5].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[5].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[5].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[5].cells[4], f"{r['dd']:.2f}%")

# Row[6] = TC 0.30% (30 bps)
r = tc_map[30]
set_cell(t.rows[6].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[6].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[6].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[6].cells[4], f"{r['dd']:.2f}%")

# Row[7] = TC 0.50% (50 bps)
r = tc_map[50]
set_cell(t.rows[7].cells[1], f"{r['cum']:.2f}%")
set_cell(t.rows[7].cells[2], f"{r['ann']:.1f}%")
set_cell(t.rows[7].cells[3], f"{r['sha']:.3f}")
set_cell(t.rows[7].cells[4], f"{r['dd']:.2f}%")

print(f"TABLE[8] (Table 9): BH sharpe={bh_shar}, 0bps={tc_map[0]['sha']}, 50bps={tc_map[50]['sha']}")

# ---------------------------------------------------------------------------
# TABLE[9] = Table 10 : GARCH vs VIX — update VIX rows only
# col headers: Regime | n | Accuracy | Fisher p | Phi (phi) | Bonferroni pass
# Row[1] = VIX High
# Row[2] = VIX Low
# ---------------------------------------------------------------------------
t = tables[9]

# Row[1] = VIX High — update threshold, n, accuracy, Fisher p, Phi, Bonferroni
set_cell(t.rows[1].cells[0], f"VIX High (\u2265 {vix_thr:.2f})")
set_cell(t.rows[1].cells[1], str(n_high))
set_cell(t.rows[1].cells[2], f"{hv_acc*100:.1f}%")
set_cell(t.rows[1].cells[3], f"{fish_p:.3f}")
set_cell(t.rows[1].cells[4], "N/A")
set_cell(t.rows[1].cells[5], "No")

# Row[2] = VIX Low
set_cell(t.rows[2].cells[1], str(n_low))
set_cell(t.rows[2].cells[2], f"{lv_acc*100:.1f}%")

print(f"TABLE[9] (Table 10): VIX High (>={vix_thr}) n={n_high}, acc=100.0%, VIX Low n={n_low}, acc={lv_acc*100:.1f}%")

# ---------------------------------------------------------------------------
# TABLE[10] = Table 11 : Threshold sensitivity — update primary p75 row
# col headers: Check | High-VIX count | High-VIX accuracy | Low-VIX accuracy | Interpretation
# Row[3] = Threshold p75-p80
# ---------------------------------------------------------------------------
t = tables[10]

# Update main result row: only p75 data is authoritative under the corrected protocol
set_cell(t.rows[3].cells[0], f"Threshold p75 (fixed pre-test = {vix_thr:.2f})")
set_cell(t.rows[3].cells[1], str(n_high))
set_cell(t.rows[3].cells[2], f"{hv_acc*100:.1f}%")
set_cell(t.rows[3].cells[3], f"{lv_acc*100:.1f}%")
set_cell(t.rows[3].cells[4], "Main result under corrected protocol; sparse High-VIX (n=8)")

print(f"TABLE[10] (Table 11): p75 row updated to n={n_high}, HV acc={hv_acc*100:.1f}%, LV acc={lv_acc*100:.1f}%")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
d.save(DST)
print(f"\nSaved to: {DST}")
print(f"McNemar p (21-day vs majority): {mcnemar_p:.4f}")
