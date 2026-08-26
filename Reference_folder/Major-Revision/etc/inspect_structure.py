import docx
import sys
import io

OUT = r"f:\MLSAPU\PhD-SPPU\India-VIX-Major-Revision\Major-Revision\etc\structure.txt"
out_f = open(OUT, "w", encoding="utf-8")
def print(*args, **kwargs):
    kwargs["file"] = out_f
    __builtins__.print(*args, **kwargs)

PATH = r"f:\MLSAPU\PhD-SPPU\India-VIX-Major-Revision\Major-Revision\Latest_Main-Manuscript_UPDATED_2026-07-24.docx"
d = docx.Document(PATH)

print("=== BODY ELEMENT SEQUENCE (paragraphs + tables in order) ===")
body = d.element.body
p_idx = 0
t_idx = 0
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        # find matching paragraph object
        text = None
        for p in d.paragraphs:
            if p._p is child:
                text = p.text
                break
        print(f"P[{p_idx}] {text[:100] if text else ''}")
        p_idx += 1
    elif tag == 'tbl':
        for t in d.tables:
            if t._tbl is child:
                nrows = len(t.rows)
                ncols = len(t.columns)
                print(f"  >>> TABLE[{t_idx}] rows={nrows} cols={ncols} header={[c.text for c in t.rows[0].cells]}")
                break
        t_idx += 1
