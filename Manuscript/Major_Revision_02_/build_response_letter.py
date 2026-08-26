import docx
from docx.shared import Pt

d = docx.Document()

def h(text, bold=True, size=12):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def para(text):
    p = d.add_paragraph(text)
    return p

def rule():
    d.add_paragraph('-' * 118)

h('Response to Reviewers', size=14)
para('Manuscript ID: NCAA-D-26-02775R2')
para('Title: India VIX as a Regime Filter for Bank Nifty Directional Prediction: A Leakage-Safe Reliability Assessment '
     '(previously: "India VIX Regime-Conditional Directional Prediction of Bank Nifty: A Stacking Ensemble Approach" -- '
     'title revised to match the corrected framing; see summary below)')
para('Journal: Neural Computing and Applications')
para('Dear Editor and Reviewers,')
para('We thank Reviewer #2 for confirming the first round of corrections and Reviewer #4 for an unusually careful, '
     'itemized reading that caught something we should have caught ourselves: the corrected analysis from the previous '
     'round had reached most of the manuscript\'s text and headline table rows, but not its figures, several table '
     'sub-rows, or two model horizons (1-day, 5-day) that were never actually retrained on the corrected split. That is '
     'a real gap, not a matter of interpretation, and we have closed it rather than defended it. Below we respond to '
     'each point.')
rule()

h('Reviewer #2')
para('Comment: The authors have put significant effort in revising and improving the article. No further comments; '
     'recommend acceptance in current form.')
para('Response: Thank you. No action required from this comment alone; the manuscript has nonetheless changed further '
     'in response to Reviewer #4, as detailed below, so Reviewer #2 may notice additional changes beyond the point that '
     'prompted their approval.')
rule()

h('Reviewer #4')

h('R4.1 Framing still argues the old case in several passages', bold=True)
para('Reviewer comment: The Introduction still states "the answer is yes, but only at the 21-trading-day horizon" and '
     'cites "market exposure on only 48 out of 277 test days"; Section 5.2 calls the regime filter "the core '
     'expert-system contribution of the paper"; Table 1 still specifies "contemporaneous India VIX"; Section 3.4 keeps '
     'a paragraph defending the threshold it has just replaced.')
para('Response: Accepted, and corrected throughout rather than in isolated spots.')
para('Changes made:')
para('- Introduction (two paragraphs): rewritten to state plainly that the answer is "not reliably, at least not in '
     'this sample," and to report the corrected trading exposure (21 of 256 aligned days, one trade, does not exceed '
     'buy-and-hold Sharpe) instead of the superseded 48/277 figure.')
para('- Section 3.4: the paragraph defending the "contemporaneous test-period threshold" as a valid diagnostic has been '
     'replaced with a paragraph describing the actual pre-test-fixed threshold methodology used throughout the '
     'corrected pipeline.')
para('- Section 5.2: the "core expert-system contribution" sentence has been replaced with an explicit statement that '
     'the pattern is a "proof-of-concept," not a validated contribution, given n=8 and non-significant tests.')
para('- Table 1: the cell reading "using contemporaneous India VIX" now reads "using a pre-test-fixed India VIX '
     'threshold (not the test-period distribution)."')
para('Location in revised manuscript: Introduction (paragraphs 2-3), Section 3.4, Section 5.2, Table 1.')

h('R4.2 No figure in Sections 5 or 6 had been rebuilt', bold=True)
para('Reviewer comment: Figure 3 still plots 0.845/0.548 with p<0.004 asterisks; Figure 5 labels its threshold line '
     '"15.08" and shades High-VIX repeatedly through Jan-Jun 2025; Figure 7 shows the strategy beating buy-and-hold at '
     'every cost level; Figure 8 is labelled n=58/n=219; Figure 9 carries "HV obs: 84.5%"; Figure 4 reports a 21-day '
     'High-VIX AUC of 0.6440, which is undefined for a single-class (n=8) subset.')
para('Response: Accepted and verified independently -- we regenerated every figure directly from the corrected result '
     'files and re-embedded them, rather than editing the old images.')
para('Changes made:')
para('- Figure 3: rebuilt as a single, honest accuracy-by-horizon chart with the correct n and accuracy values and an '
     'explicit non-significance caption.')
para('- Figure 4: rebuilt to show only defensible AUC values; the High-VIX 21-day cell is now marked "N/A (single '
     'class)" rather than a fabricated-looking 0.644.')
para('- Figure 5: rebuilt from the actual corrected trade simulation (reconstructed day-by-day from '
     'test_predictions.csv and raw OHLCV, matching the reported summary statistics exactly): threshold line now reads '
     '"18.71," and the shaded region shows the single real trade window (2025-04-08 to 2025-05-12), not repeated '
     'shading across Jan-Jun 2025.')
para('- Figure 6: rebuilt as a proper drawdown time series from the same reconstruction (max DD -3.69% vs -6.62%, '
     'matching Table 8).')
para('- Figure 7: replaced with the actual corrected transaction-cost-sensitivity plot (Sharpe declining from 0.455 to '
     '0.375, below the buy-and-hold line at 1.204 throughout).')
para('- Figure 8: rebuilt from a newly performed, leakage-safe GARCH(1,1) analysis (see R4.3) with correct n values.')
para('- Figure 9: replaced with the corrected bootstrap/permutation diagnostic plot (95% CI [16.2, 51.4] pp, observed '
     '33.5 pp marked against both the bootstrap and permutation null distributions).')
para('Location in revised manuscript: Figures 3-9 (Sections 5.2-6.3).')

h('R4.3 Table 10 internally impossible (intersection exceeds parent set; inconsistent sample sizes)', bold=True)
para('Reviewer comment: The VIX-High/GARCH-High intersection is reported as 47 while VIX-High itself has only 8 '
     'members; VIX rows total 256 while GARCH rows total 277.')
para('Response: Accepted. On investigation, the GARCH comparison in the previous submission traced back to a notebook '
     '(external to this manuscript\'s corrected pipeline) that itself fits GARCH(1,1) on train+test combined data -- '
     'the same look-ahead-bias category this revision otherwise corrects -- and contains incomplete/non-executing code. '
     'We did not attempt to patch that analysis; we performed a new one.')
para('Changes made:')
para('- A new GARCH(1,1) regime analysis was implemented: fit on train+validation returns only (pre-test), with the '
     'GARCH-High threshold set at the 75th percentile of the pre-test conditional volatility, and out-of-sample '
     'conditional volatility for the test period produced by expanding day-by-day re-estimation (no future '
     'information used at any point).')
para('- The comparison is now computed on the identical 256-row aligned test sample used everywhere else in the paper. '
     'VIX-High n=8 (100.0% accuracy), GARCH-High n=10 (80.0% accuracy), intersection n=2 (both non-overlapping Fisher '
     'p=1.000). The intersection can no longer exceed either parent set by construction.')
para('- Section 6.2 no longer claims India VIX is "the stronger gatekeeper": under the corrected protocol, neither '
     'regime split is statistically significant, and we say so explicitly.')
para('Location in revised manuscript: Section 6.2, Table 10, Figure 8.')

h('R4.4 Table 11 internally impossible (p85 selects more days than p75; quarterly rows sum to the superseded count)', bold=True)
para('Reviewer comment: p75 selects 8 days and p85 selects 28, which no single distribution allows; the 2025-Q1 and Q2 '
     'rows (17 and 41) sum to 58, the count from the original, biased threshold.')
para('Response: Accepted and recomputed from the corrected pre-test VIX distribution.')
para('Changes made:')
para('- Threshold sensitivity now runs monotonically as required: p60=31, p65=27, p70=16, p75=8, p85=4 High-VIX days, '
     'each computed from the train+validation VIX distribution and applied to the same 256-row test sample.')
para('- The quarterly breakdown is recomputed under the corrected p75 threshold: all 8 High-VIX days fall in 2025-Q2, '
     'summing correctly to 8 (not 58).')
para('- The "P1 replication" row has been removed. On inspection, the only available secondary pipeline (P1) computes '
     'its own regime threshold from the test-period VIX distribution (the same look-ahead-bias pattern this revision '
     'corrects for the primary pipeline), and different snapshots of its regime split in the underlying project '
     'directory are mutually inconsistent in both window and High-VIX count -- none of which we discovered until '
     'checking the source data for this response. We report this transparently rather than retain a number we cannot '
     'stand behind; an independently trained, pre-test-threshold replication is noted as a concrete next step '
     '(Section 6.5, Section 7).')
para('Location in revised manuscript: Section 6.3, Table 11, Section 6.5.')

h('R4.5 Shorter horizons (1-day, 5-day) unchanged from the pre-correction version', bold=True)
para('Reviewer comment: Table 6\'s 1-day and 5-day rows are unchanged to four decimal places from the previous version, '
     'and Table 7\'s High-VIX counts (58, 57) look untouched by the ~21-row split correction and full retrain.')
para('Response: Confirmed and corrected. On investigation, this suspicion was entirely right: only the 21-day model '
     'was ever retrained on the corrected split in the previous round. The 1-day and 5-day rows were leftover from the '
     'original, uncorrected submission. We have now trained real 1-day and 5-day stacking ensembles -- same '
     'architecture, same corrected splits, same pre-test-fixed regime threshold -- and replaced every 1-day/5-day '
     'number in the manuscript with the output of that run.')
para('Changes made:')
para('- New stacking ensembles trained for both horizons (XGBoost, LightGBM, Random Forest, CatBoost, BiLSTM with '
     'self-attention, ridge meta-learner), using the same corrected train/validation/test splits and the same 16 '
     'leakage-safe features. Optuna trials were reduced from 100 to 25 per model for turnaround time; this is stated '
     'explicitly rather than left implicit.')
para('[[HORIZON_RESULTS_PLACEHOLDER]]')
para('Location in revised manuscript: Table 5, Table 6, Section 5.1, Introduction.')

h('R4.6 Sharpe conventions not stated; single-trade Sharpe questioned as a distribution statistic', bold=True)
para('Reviewer comment: Section 5.3 gives execution timing but not the risk-free rate or annualisation basis; a Sharpe '
     'ratio built from one position across 256 mostly-flat days may not be a meaningful distribution statistic.')
para('Response: Accepted on the first point; the underlying computation was always well-defined but was not stated in '
     'the manuscript text, and the Response letter for the previous round incorrectly told Reviewer #2 the risk-free '
     'rate was 0% when the code actually used 6.5%. Both are now fixed.')
para('Changes made:')
para('- Section 5.3 now states explicitly: risk-free rate 6.5% p.a. (India 91-day T-bill, RBI FY2025 average), daily '
     'RFR = 6.5%/252, Sharpe annualised as (mean daily excess return / std daily excess return) x sqrt(252), computed '
     'identically for the strategy and buy-and-hold series over the same 256-day window.')
para('- On the second point, we agree the resulting Sharpe (one position, 235 of 256 days flat) is a weak statistic '
     'and now say so in Section 6.1 rather than presenting it as decisive; the raw trade-level return (single 21-day '
     'trade, +9.91% before costs) is also reported.')
para('Location in revised manuscript: Section 5.3, Section 6.1.')

h('R4.7 GARCH split reported as the only significant regime effect, contradicting the manuscript\'s own numbers', bold=True)
para('Reviewer comment: Section 6.2 concludes India VIX is "the stronger gatekeeper," but Table 10 shows the GARCH '
     'split as nominally significant (p=0.039) against p=1.000 for VIX.')
para('Response: This was correct as a critique of the previous submission\'s Table 10, but is now moot because that '
     'GARCH analysis has been replaced (R4.3). Under the corrected, leakage-safe GARCH analysis, neither split is '
     'significant (both non-overlapping Fisher p=1.000), so the "stronger gatekeeper" claim has been removed entirely '
     'rather than reworded to favor either variable.')
para('Location in revised manuscript: Section 6.2.')

h('R4.8 Cost-sensitivity and bootstrap sections state conclusions the surrounding numbers do not support', bold=True)
para('Reviewer comment: Section 6.1 says the strategy pattern is "not" one that only looks good before costs, in a '
     'passage showing it below buy-and-hold at every cost level; Section 6.3 reports a bootstrap interval excluding '
     'zero alongside non-significant tests without resolving the tension; the permutation p-value was not printed.')
para('Response: Accepted.')
para('Changes made:')
para('- Section 6.1 no longer claims the strategy escapes the "looked good only before costs" pattern; it states '
     'plainly that the strategy is below buy-and-hold at every tested cost level, and that the constant drawdown '
     'reflects there being exactly one position, not risk-selection skill.')
para('- Section 6.3 now prints the block-permutation p-value (0.357 before Bonferroni, 1.000 after) alongside the '
     'bootstrap CI, and states directly that a bootstrap interval excluding zero and a non-significant permutation '
     'test can coexist under small, non-i.i.d. samples -- this is a caution about the sample, not a resolved finding '
     'in either direction.')
para('Location in revised manuscript: Section 6.1, Section 6.3.')

h('R4.9 Smaller inconsistencies (naive-majority rate, old-threshold day count, GNN references, uncited bootstrap procedure)', bold=True)
para('Reviewer comment: Naive-majority rate given as 71.1% in Table 6 and 68.9% in old Figures 8/9; Section 6.3 gives '
     'the old threshold as yielding 56 High-VIX days where the letter and Table 7 use 58; block bootstrap/permutation '
     'procedures uncited.')
para('Response: Accepted.')
para('Changes made:')
para('- All regenerated figures use the same majority-baseline convention as the tables (71.1% overall); the old '
     'figures carrying 68.9% have been replaced (R4.2).')
para('- The original-threshold High-VIX day count is now stated consistently as 56 throughout (matching the '
     'vix_threshold_config.json artifact), and the Response letter\'s "58" is acknowledged as the less reliable figure.')
para('- We were not able to locate a citable source paper specific to the exact circular block-bootstrap/permutation '
     'implementation used; the procedures are standard (Politis & Romano-style block resampling; permutation testing '
     'under label exchangeability), and Section 6.3 now describes the procedure parameters (block length 21, n=10,000 '
     'resamples) in enough detail to be independently reproduced, in lieu of a single canonical citation.')
para('Location in revised manuscript: Figures 3-9, Section 6.3.')

h('On the reviewer\'s closing suggestion', bold=True)
para('Reviewer comment: "Two ordinary mistakes accounted for the whole of the original result... A paper built around '
     'that, with the walk-forward at its centre and a window covering more than one stress episode, would be worth '
     'reading... It would need a different title."')
para('Response: We have taken this suggestion directly rather than treating it as optional framing advice.')
para('Changes made:')
para('- A new Section 6.4 (Walk-Forward Generalizability) has been added, reporting year-by-year regime accuracy '
     'across six independently evaluated years (2019, 2020, 2021, 2022, 2024, 2025; 2023 is the validation-boundary '
     'year). The High-VIX advantage holds in exactly one year (2021, +12.4 pp) and reverses in the other five (-11.6 '
     'to -39.3 pp). This is now the paper\'s central robustness result and is stated as such in the Abstract, the '
     'Introduction, and the Conclusion.')
para('- The Abstract and Conclusion have been rewritten to lead with this finding: the paper is now positioned as a '
     'methodological case study in how a test-period-derived threshold and an overlapping-position trading simulation '
     'combined to manufacture the original submission\'s positive result, together with a fully corrected '
     'reliability-assessment framework -- not as a validated trading edge.')
para('- The title has been changed to "India VIX as a Regime Filter for Bank Nifty Directional Prediction: A '
     'Leakage-Safe Reliability Assessment," reflecting this reframing while retaining enough continuity with the '
     'original title for the editorial record.')
para('- We have not pursued a genuinely independent multi-episode replication window in this round -- doing so '
     'properly is exactly the "new work" the reviewer correctly identified it as, and the walk-forward analysis '
     'above is, we believe, a reasonable interim substitute using data already available within the study\'s own '
     'window. We would welcome the opportunity to pursue a true multi-year replication pipeline in a future revision '
     'if the Editor judges it necessary.')
para('Location in revised manuscript: Abstract, Introduction, Section 6.4, Conclusion, Title.')

h('On the reviewer\'s request for analysis code', bold=True)
para('Reviewer comment: "I would welcome access to the analysis code, in particular the split construction, the '
     'threshold calculation and the trading simulation."')
para('Response: We can provide the notebooks implementing the split construction (NB-R01), threshold calibration '
     '(NB-R02), model training (NB-R03, and the newly added NB-R13 for the 1-day/5-day horizons), statistical testing '
     '(NB-R04), walk-forward analysis (NB-R05), the GARCH and extended threshold-sensitivity analysis added in this '
     'round (NB-R11), figure regeneration (NB-R12), and the trading simulation (NB-R09) as supplementary material, or '
     'via a code-availability statement, at the Editor\'s discretion.')

rule()
h('Summary of Revision Impact (this round)')
para('- Every figure in Sections 5-6 (Figs 3-9) rebuilt from the corrected pipeline and verified against the '
     'underlying result files; no figure now shows pre-correction numbers.')
para('- Table 10 (VIX vs GARCH) and Table 11 (threshold sensitivity) recomputed from a new, leakage-safe GARCH(1,1) '
     'analysis and a corrected threshold sweep; both internal impossibilities (intersection > parent set; p85 count > '
     'p75 count) are resolved by construction.')
para('- 1-day and 5-day models retrained for the first time on the corrected, leakage-safe split (previously never '
     'retrained after the original submission).')
para('- Missing analyses actually merged into the manuscript body: SHAP results (Section 5.4, new Fig. 11), ablation '
     'study (new Section 6.6, Table 13, Fig. 13), walk-forward generalizability (new Section 6.4, Table 12, Fig. 12), '
     'regime-stratified baseline comparison (Section 5.1, Table 6b, Fig. 3b), and real tuned hyperparameters (Table 4).')
para('- Sharpe convention (6.5% RFR, 252-day annualisation) stated explicitly in-text; the previous round\'s incorrect '
     '"0%" claim to Reviewer #2 corrected.')
para('- Eight passages of leftover pre-correction framing rewritten (Introduction x2, Section 3.4, Section 5.2, '
     'Section 5.3, Section 6.2 x2, Section 6.5), plus Table 1\'s wording and the removal of an unreliable '
     'cross-pipeline comparison figure.')
para('- Abstract, Conclusion, and Title reframed around the walk-forward non-replication finding as the paper\'s '
     'central, honest contribution, per the reviewer\'s explicit suggestion.')
para('')
para('We thank Reviewer #4 again for an exceptionally thorough and fair reading. The gap between what we believed we '
     'had delivered and what had actually reached the manuscript file was real, and this round of revision has been '
     'about closing that gap rather than arguing against it.')
para('Sincerely,')
para('All Authors')

d.save('Response_to_Reviewers_R2.docx')
print('Draft response letter saved.')
